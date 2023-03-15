import math
import os
import sys
from marko.renderer import Renderer
from marko import Markdown
from marko.helpers import camel_to_snake_case

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

import logging
import argparse
import html

_FONT_NAME = None
_FONT_COLOR = RGBColor(0, 0, 0)
_FONT_SIZE = Pt(12)
_FONT_BOLD = False

_PARAGRAPH_SPACING = Pt(0.0)
_PARAGRAPH_LINE_SPACING = 2.0
_PARAGRAPH_WIDOW_CONTROL = True

_MARGIN = Inches(1.0)

_BUFFER_PARAGRAPHS = 7

_STYLE_NAMES = [
    "Normal",
    "Body Text", "Body Text 2", "Body Text 3",
    "Heading 1", "Heading 2",
]

_DOCUMENT = None
_VARIABLES = None

_DEFAULT_CONTACT = [
    "Contact Name",
    "476 5th Avenue",
    "New York, NY 10018",
    "(212) 555-1234",
    "email@example.com",
    "",
    "Organizational memberships go here"
]

class __DocRenderer(Renderer):
    def __init__(self):
        # Our core document
        self.document = _DOCUMENT

        self.variables = _VARIABLES

        # A stack of custom styles for rendering raw text
        self.style = [None]

        self._init_styles()
        self._init_margins()
        self._init_header()
        self._add_buffer_paragraphs()
        self._add_title()

    def _init_styles(self):
        for style_name in _STYLE_NAMES:
            style = self.document.styles[style_name]

            style.font.name = _FONT_NAME
            style.font.size = _FONT_SIZE
            style.font.color.rgb = _FONT_COLOR
            style.font.bold = _FONT_BOLD
            style.paragraph_format.space_before = _PARAGRAPH_SPACING
            style.paragraph_format.space_after = _PARAGRAPH_SPACING

    def _init_margins(self):
        section = self.document.sections[0]
        section.left_margin = _MARGIN
        section.right_margin = _MARGIN
        section.top_margin = _MARGIN
        section.bottom_margin = _MARGIN

    def _init_header(self):
        section = self.document.sections[0]
        section.different_first_page_header_footer = True

        # First page headers
        p = section.first_page_header.paragraphs[0]

        contact = self.variables.get("contact", [""])
        run = p.add_run("{}\t\tabout {} words".format(
            contact[0], self.variables.get("word_count")))
        _set_style(run)
        for i in range(1, len(contact)):
            run = p.add_run("\n{}".format(contact[i]))
            _set_style(run)

        # Other page headers
        p = section.header.paragraphs[0]
        run = p.add_run(
            "\t\t{} / {} / ".format(
                self.variables.get("header_author"), self.variables.get("header_title")
            )
        )
        _set_style(run)

        _add_field(run, "PAGE")

        run = p.add_run("\n\n")
        _set_style(run)

    def _add_buffer_paragraphs(self):
        for _ in range(_BUFFER_PARAGRAPHS):
            p = self.document.add_paragraph()
            p.paragraph_format.line_spacing = _PARAGRAPH_LINE_SPACING
            p.paragraph_format.widow_control = _PARAGRAPH_WIDOW_CONTROL

    def _add_title(self):
        h = self.document.add_heading()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run("{}\n\nby {}\n\n".format(
            self.variables.get("title"), self.variables.get("author")
        ))
        _set_style(run)

    def render_children(self, element):
        # If sent a list, each one is a piece of text to be rendered with
        # no children. Render them and return.
        if isinstance(element, list):
            for e in element:
                self.render_children(e)

            return

        # Figure out the type of element and extract its potential children
        # safely.
        element_class = camel_to_snake_case(element.__class__.__name__)
        if hasattr(element, "children"):
            children = element.children
        else:
            children = []

        # Add a new paragraph
        if element_class == "paragraph":
            p = self.document.add_paragraph()
            p.paragraph_format.line_spacing = _PARAGRAPH_LINE_SPACING
            p.paragraph_format.widow_control = _PARAGRAPH_WIDOW_CONTROL

            run = p.add_run("\t")

            self.render_children(children)
            return

        # Add a new heading
        elif element_class == "heading":
            h = self.document.add_heading()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self.render_children(children)
            return

        # Render raw text within the most recent paragraph, potentially
        # with emphasis.
        elif element_class == "raw_text":
            if len(self.document.paragraphs) < 1:
                p = self.document.add_paragraph()
                p.paragraph_format.line_spacing = _PARAGRAPH_LINE_SPACING
                p.paragraph_format.widow_control = _PARAGRAPH_WIDOW_CONTROL

            string = html.unescape(children)
            p = self.document.paragraphs[-1]

            if self.style[-1] == "emphasis":
                run = p.add_run(string)
                _set_style(run)
                run.italic = True
            else:
                run = p.add_run(string)
                _set_style(run)

            return

        # Start rendering with emphasis, then pop the style from the stack.
        elif element_class == "emphasis":
            self.style.append("emphasis")
            self.render_children(children)
            self.style.pop()

            return

        # In the default case we ignore the element but still walk its
        # children.
        logging.debug("Ignoring: {}".format(element))

        self.render_children(children)
        return

def _set_style(run):
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE
    run.font.color.rgb = _FONT_COLOR
    run.font.bold = _FONT_BOLD

## Stolen from https://github.com/python-openxml/python-docx/issues/498
def _add_field(run, field):
    fldChar1 = OxmlElement("w:fldChar")  # creates a new element
    fldChar1.set(qn("w:fldCharType"), "begin")  # sets attribute on element
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")  # sets attribute on element
    instrText.text = field

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click to update field."
    fldChar2.append(t)

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

def _replace_quotes(content_str):
    content = list(content_str)
    left_quote = "“"
    right_quote = "”"
    apostrophe = "’"
    q = 0
    for i in range(len(content)):
        match content[i]:
            case "'":
                content[i] = apostrophe
            case '"':
                if q % 2 == 0:
                    content[i] = left_quote
                else:
                    content[i] = right_quote
                q += 1

    if q % 2 == 1:
        raise Exception("Non-even number of quotation marks")

    return "".join(content)

def _default_header_title(title):
    title_words = title.split()
    for i in range(len(title_words)):
        if len(title_words[i]) > 3:
            return title_words[i]

    return title_words[0]

def _default_header_author(author):
    return author.split()[-1]

def _compute_output_file(args):
    if args.output is not None:
        return args.output

    elif args.input == '-':
        return sys.stdout

    else:
        input_wo_ext, _ = os.path.splitext(args.input)
        return input_wo_ext + ".docx"

def main():
    # Workarounds for custom Markdown renderer.
    global _VARIABLES, _DOCUMENT, _FONT_NAME

    md = Markdown(renderer=__DocRenderer)

    parser = argparse.ArgumentParser(description="md2shunn")
    parser.add_argument(
        "--input", default=None, type=str, required=True, help="Markdown input file. Use '-' for stdin",
    )
    parser.add_argument(
        "--output", default=None, type=str,
        help="Word .docx output file. Defaults to the same base filename of the input file but with a .docx " +
            "extension unless the input file is stdin, in which case the output is stdout. Use '-' to explicitly " +
            "set stdout"
    )
    parser.add_argument(
        "--title", default="TITLE", type=str, help="Title of the piece. Defaults to placeholder value.",
    )
    parser.add_argument(
        "--author", default="AUTHOR", type=str, help="Author of the piece. Defaults to placeholder value.",
    )
    parser.add_argument(
        "--header-title", default=None, type=str,
        help="Short version of the title to use in header. Defaults to an auto-truncated version of --title",
    )
    parser.add_argument(
        "--header-author", default=None, type=str,
        help="Last name of the author to use in header. Defaults to an auto-truncated version of --author",
    )
    parser.add_argument(
        "--format", default="modern", type=str, choices=["modern", "classic"],
        help="Use 'modern' (variable-width) or 'classic' (fixed-width) font. Defaults to 'modern'",
    )

    args = parser.parse_args()

    if args.header_author is None:
        args.header_author = _default_header_author(args.author)

    if args.header_title is None:
        args.header_title = _default_header_title(args.title)

    args.contact = _DEFAULT_CONTACT

    match args.format:
        case "modern":
            _FONT_NAME = "Times New Roman"
        case "classic":
            _FONT_NAME = "Courier New"
        case _:
            raise Exception("Unexpected format: " + args.format)

    if args.input == '-':
        content = sys.stdin.read()
    else:
        with open(args.input, "r") as f:
            content = _replace_quotes(f.read())

    word_count = int(math.ceil(len(content.split()) / 100) * 100)
    word_count_formatted = f'{word_count:,}'

    _DOCUMENT = Document()
    _VARIABLES = {
        "title": args.title,
        "author": args.author,
        "header_title": args.header_title,
        "header_author": args.header_author,
        "word_count": word_count_formatted,
        "contact": _DEFAULT_CONTACT
    }

    md(content)
    output_file = _compute_output_file(args)
    if output_file == sys.stdout:
        with os.fdopen(sys.stdout.fileno(), "wb", closefd=False) as out:
            _DOCUMENT.save(out)
    else:
        _DOCUMENT.save(output_file)

if __name__ == "__main__":
    main()
